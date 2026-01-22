from docx import Document
from docx.shared import Pt
import re
import os

def markdown_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"File {md_file} not found.")
        return

    with open(md_file, 'r') as f:
        md_content = f.read()

    doc = Document()

    # Simple Markdown Parser
    lines = md_content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph()
            # Handle bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file)
    print(f"Generated {docx_file}")

if __name__ == "__main__":
    markdown_to_docx("ANALYTICAL_NOTES.md", "ANALYTICAL_NOTES.docx")
